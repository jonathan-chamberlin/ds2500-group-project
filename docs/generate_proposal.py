from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import re

doc = Document()

# set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
pf = style.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
pf.space_after = Pt(0)
pf.space_before = Pt(0)

# set margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('DS 2500 Project Proposal')
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'
title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

# team info
team_info = doc.add_paragraph()
team_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
team_info.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
lines = [
    'Jonathan Chamberlin (chamberlin.j@northeastern.edu)',
    'Min Yu Huang (huang.minyu@northeastern.edu)',
    'Anuhya Mandava (mandava.an@northeastern.edu)',
    'Tsion Teklaeb (teklaeb.t@northeastern.edu)',
    'Section 1'
]
for i, line in enumerate(lines):
    run = team_info.add_run(line)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if i < len(lines) - 1:
        team_info.add_run('\n')


def add_heading_text(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def add_body_text(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


# problem statement
add_heading_text(doc, 'Problem Statement')

add_body_text(doc, (
    "We want to figure out which state-level health and lifestyle factors are the "
    "best predictors of chronic disease rates across the US. Specifically, we are "
    "looking at state-level prevalence rates (percentage of population affected) for "
    "diabetes, heart disease, and COPD as our target outcomes, and testing whether "
    "indicators like smoking rates, obesity, alcohol use, household income, and "
    "insurance coverage can predict how those rates vary from state to state. Some "
    "states have way higher chronic disease rates than others, and we want to "
    "understand what is actually driving those gaps."
))

add_body_text(doc, (
    "This is scoped to state-level analysis using existing public datasets, which "
    "keeps it manageable for one semester and still gives us enough variation across "
    "50+ states to build meaningful models."
))

# data sources
add_heading_text(doc, 'Data Sources')

add_body_text(doc, (
    "Our primary dataset is the U.S. Chronic Disease Indicators dataset published "
    "by the CDC (https://catalog.data.gov/dataset/u-s-chronic-disease-indicators). "
    "It covers 115 health indicators tracked across all US states and territories, "
    "including diabetes, COPD, tobacco use, cardiovascular disease, and more. The "
    "data is available in CSV format and is organized at the state level."
))

add_body_text(doc, (
    "We also plan to pull in supplementary data from the National Alcohol Surveys "
    "(https://arg.org/center/national-alcohol-surveys/) for more detailed "
    "alcohol-related risk factors, and potentially US Census data for socioeconomic "
    "variables like household income and insurance coverage rates."
))

add_body_text(doc, (
    "Ethical Considerations: Since all of our data is aggregated at the state level, "
    "there is no risk of identifying individual people. But there are still things to "
    "be careful about. State-level averages can hide disparities within a state, so "
    "our conclusions might not apply equally to all communities. Some populations are "
    "underrepresented in health surveys due to barriers like language or healthcare "
    "access, which means the data could undercount certain groups. We also need to be "
    "thoughtful about how we frame our results, since labeling states as \"high risk\" "
    "could affect how resources get allocated or reinforce existing stereotypes about "
    "certain regions. We will frame our findings around systemic factors instead of "
    "blaming individual states, and note limitations of the data where relevant."
))

# analysis goals and methods
add_heading_text(doc, 'Analysis Goals and Methods')

add_body_text(doc, (
    "We would like to use correlation analysis (Pearson and Spearman) and hypothesis "
    "testing to identify which health and demographic indicators have the strongest "
    "association with chronic disease prevalence across states. For example, we want "
    "to see if smoking rates or obesity levels are more predictive of diabetes rates "
    "than income or insurance coverage."
))

add_body_text(doc, (
    "We would like to build machine learning models (KNN and linear regression) to "
    "predict chronic disease rates in a given state based on its risk factor profile. "
    "We will build separate models for each disease and evaluate them using metrics "
    "like R-squared and RMSE with cross-validation. We plan to compare how well "
    "these models perform across different diseases to see whether the same factors "
    "drive diabetes, heart disease, and COPD or if each has its own key predictors."
))

add_body_text(doc, (
    "We also want to create geographic visualizations like choropleth maps and "
    "scatter plots to show spatial patterns in chronic disease burden and highlight "
    "which regions are most affected. By the end of the project, we want to be able "
    "to say which handful of factors matter the most for each disease and see if it "
    "actually works well enough to be useful for predicting state-level outcomes."
))

# division of labor
add_heading_text(doc, 'Division of Labor')

add_body_text(doc, (
    "Jonathan Chamberlin will handle data acquisition and pipeline development, "
    "including downloading, cleaning, and merging the CDC dataset with any "
    "supplementary sources. He has a programming background and will make sure the "
    "data infrastructure is set up for the rest of the team to build on. He will "
    "also contribute analysis focused on geographic trends in chronic disease."
))

add_body_text(doc, (
    "Min Yu Huang will lead the statistical analysis, running correlation tests and "
    "hypothesis testing to determine which risk factors have the strongest "
    "relationships with disease outcomes. She has experience with statistical methods "
    "and will focus specifically on diabetes prevalence as her primary analysis "
    "question."
))

add_body_text(doc, (
    "Anuhya Mandava will lead the machine learning modeling effort, building and "
    "evaluating KNN and regression models to predict chronic disease rates. She has "
    "coursework in data science and will compare model accuracy across different "
    "diseases and work on tuning parameters."
))

add_body_text(doc, (
    "Tsion Teklaeb will lead the visualization work, creating choropleth maps, "
    "trend plots, and other visual summaries of the data. She has a background in "
    "public health topics and will also take the lead on the ethical considerations "
    "section of the final report, examining biases in the data and what our findings "
    "mean for policy."
))

add_body_text(doc, (
    "All team members will contribute to writing the final report and preparing "
    "the presentation."
))

# save
output_path = 'DS2500_project_proposal_final.docx'
doc.save(output_path)
print(f"Saved to {output_path}")

# count words
all_text = ' '.join(p.text for p in doc.paragraphs)
word_count = len(re.findall(r'\w+', all_text))
print(f"Word count: {word_count}")
